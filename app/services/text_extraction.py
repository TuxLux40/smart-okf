"""Text extraction helpers for ingest pipeline.

Scanned PDFs (no text layer) are OCRed **in place** with OCRmyPDF before extraction:
the text layer is embedded into the original PDF (searchable/editable ever after) and
never needs re-running — post-OCR the PDF has text, so the OCR branch is skipped on
later ingests. Standalone images are OCRed read-only via tesseract (images have no
embeddable text layer, so their text lives only in the transcript/aggregate).

PDF extraction defaults to an external `marker` CLI backend (layout-aware: tables, forms)
— **not** a pip dependency of this project; onboarding installs `marker-pdf` externally
(`marker_single` on PATH), same pattern as `tesseract`/`ghostscript`. GPL-3.0 code +
modified OpenRAIL-M weights stay outside this MIT graph. Opt out via
`ExtractionOptions(use_marker=False)` / CLI `--no-marker` for pdfplumber + in-place
OCRmyPDF. See README.md.
"""

import os
import shutil
import subprocess
import tempfile
from email import message_from_bytes
from email.message import Message
from pathlib import Path

import openpyxl
import pdfplumber
from docx import Document as DocxDocument

from app.constants import IMAGE_DOCUMENT_SUFFIXES, OCR_LANGUAGES, SUPPORTED_DOCUMENT_SUFFIXES, TEXT_FILE_ENCODING
from app.exceptions import DocumentIngestError, EncryptedDocumentError
from app.services.extraction_options import DEFAULT_EXTRACTION, ExtractionOptions

_OLE_MAGIC = b"\xd0\xcf\x11\xe0"
"""OLE2 compound-file header. Encrypted OOXML (.docx/.xlsx) is wrapped in this container
instead of a ZIP, so a docx/xlsx starting with these bytes is password-protected."""


def is_supported_document(file_path: Path) -> bool:
    """Return whether a file suffix is supported for ingest."""
    return file_path.suffix.lower() in SUPPORTED_DOCUMENT_SUFFIXES


def _raise_if_encrypted(file_path: Path) -> None:
    """Raise `EncryptedDocumentError` for a password-protected PDF or OOXML file.

    Cheap header/parser probe run before extraction so a protected file is skipped with
    a clear reason instead of failing deep inside marker/pdfplumber/openpyxl with an
    opaque traceback. Only detects the common personal-document cases (encrypted PDF,
    encrypted .docx/.xlsx); anything else falls through to normal extraction.
    """
    suffix = file_path.suffix.lower()
    if suffix in {".docx", ".xlsx"} and file_path.read_bytes()[:4] == _OLE_MAGIC:
        raise EncryptedDocumentError(f"{file_path} is password-protected (encrypted OOXML); skipped")
    if suffix == ".pdf":
        from pdfminer.pdfdocument import PDFDocument, PDFPasswordIncorrect
        from pdfminer.pdfparser import PDFParser

        with file_path.open("rb") as handle:
            try:
                PDFDocument(PDFParser(handle))
            except PDFPasswordIncorrect as error:
                raise EncryptedDocumentError(f"{file_path} is password-protected (encrypted PDF); skipped") from error
            except Exception:  # noqa: BLE001 — not an encryption signal; let real extraction report it
                return


def extract_text_from_file(
    file_path: Path,
    options: ExtractionOptions = DEFAULT_EXTRACTION,
) -> str:
    """Extract text from a supported document.

    With `options.use_marker` (default True), `.pdf` extraction goes through the marker CLI
    backend for layout-aware extraction (tables, forms). Onboarding installs `marker_single`
    as a prerequisite; explicit-fail rather than silent fallback if missing. Pass
    `ExtractionOptions(use_marker=False)` (CLI: `--no-marker`) for pdfplumber/OCRmyPDF only.
    """
    suffix = file_path.suffix.lower()
    _raise_if_encrypted(file_path)
    if suffix in IMAGE_DOCUMENT_SUFFIXES:
        return _extract_text_from_image(file_path)
    if suffix == ".pdf":
        if options.use_marker:
            if not marker_available():
                raise DocumentIngestError("marker_single not found on PATH; install marker-pdf or pass --no-marker")
            return _extract_text_from_pdf_via_marker(file_path)
        return _extract_text_from_pdf(file_path, allow_ocr_rewrite=options.allow_ocr_rewrite)
    if suffix == ".docx":
        return _extract_text_from_docx(file_path)
    if suffix == ".eml":
        return _extract_text_from_eml(file_path)
    if suffix == ".xlsx":
        return _extract_text_from_xlsx(file_path)
    return file_path.read_text(encoding=TEXT_FILE_ENCODING, errors="ignore")


def _extract_text_from_pdf(file_path: Path, *, allow_ocr_rewrite: bool = True) -> str:
    """Extract plain text from a PDF; OCR scanned PDFs in place first (when permitted).

    With `allow_ocr_rewrite=False` (read-only paths like transcript backfill), a scanned
    PDF raises instead of being silently mutated — the caller treats it as best-effort.
    """
    text = _read_pdf_text(file_path)
    if text.strip():
        return text
    if not allow_ocr_rewrite:
        raise DocumentIngestError(
            f"scanned PDF {file_path} needs OCR, but this pass is read-only (no in-place rewrite)"
        )
    ocr_pdf_in_place(file_path)
    return _read_pdf_text(file_path)


def _read_pdf_text(file_path: Path) -> str:
    """Extract the existing text layer from a PDF using pdfplumber."""
    pages: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                pages.append(page_text)
    return "\n".join(pages)


def ocr_pdf_in_place(file_path: Path) -> None:
    """Embed an OCR text layer into a scanned PDF, replacing the original atomically.

    Uses OCRmyPDF with `--skip-text` so pages that already carry text are untouched.
    The embedded layer persists in the PDF itself — later ingests and PDF editors get
    the text for free.
    """
    import ocrmypdf

    with tempfile.NamedTemporaryFile(suffix=".pdf", dir=file_path.parent, delete=False) as temp:
        temp_path = Path(temp.name)
    try:
        ocrmypdf.ocr(
            file_path,
            temp_path,
            language=OCR_LANGUAGES,
            skip_text=True,
            progress_bar=False,
        )
        os.replace(temp_path, file_path)
    except Exception as error:
        temp_path.unlink(missing_ok=True)
        raise DocumentIngestError(f"OCR failed for {file_path}") from error


def marker_available() -> bool:
    """Whether a separately-installed `marker_single` binary is on PATH.

    Public so CLI entry points can preflight once and fail fast with a single clear
    error, instead of every PDF being individually skipped in a "successful" run.
    """
    return shutil.which("marker_single") is not None


def _marker_error_detail(error: BaseException) -> str:
    """Best-effort stderr/stdout snippet for marker failures (never empty noise only)."""
    if isinstance(error, subprocess.TimeoutExpired):
        return "timed out after 600s"
    if isinstance(error, subprocess.CalledProcessError):
        parts = [f"exit {error.returncode}"]
        for label, stream in (("stderr", error.stderr), ("stdout", error.stdout)):
            if isinstance(stream, str) and stream.strip():
                parts.append(f"{label}: {stream.strip()[-2000:]}")
            elif isinstance(stream, bytes) and stream.strip():
                parts.append(f"{label}: {stream.strip()[-2000:].decode('utf-8', errors='replace')}")
        return "; ".join(parts)
    return str(error)


def _extract_text_from_pdf_via_marker(file_path: Path) -> str:
    """Shell out to a user-installed `marker_single` CLI for layout-aware PDF extraction.

    Marker is never a pip dependency of smart-okf — onboarding installs `marker-pdf`
    externally (pipx/sibling venv). This path does not run OCRmyPDF in-place; marker
    handles layout + any OCR it needs itself. Use `ExtractionOptions(use_marker=False)`
    for the pdfplumber + in-place OCRmyPDF pipeline (searchable PDF rewrite).
    """
    with tempfile.TemporaryDirectory() as out_dir:
        try:
            subprocess.run(
                ["marker_single", str(file_path), "--output_dir", out_dir, "--output_format", "markdown"],
                capture_output=True,
                text=True,
                check=True,
                timeout=600,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as error:
            detail = _marker_error_detail(error)
            raise DocumentIngestError(f"marker extraction failed for {file_path}: {detail}") from error
        produced = next(Path(out_dir).rglob("*.md"), None)
        if produced is None:
            raise DocumentIngestError(f"marker produced no output for {file_path}")
        return produced.read_text(encoding="utf-8")


def _extract_text_from_image(file_path: Path) -> str:
    """OCR a standalone image via tesseract (read-only; the image is not modified)."""
    try:
        completed = subprocess.run(
            ["tesseract", str(file_path), "stdout", "-l", OCR_LANGUAGES],
            capture_output=True,
            text=True,
            errors="replace",
            check=True,
            timeout=300,
        )
    except FileNotFoundError as error:
        raise DocumentIngestError(f"tesseract not installed; cannot OCR {file_path}") from error
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as error:
        raise DocumentIngestError(f"OCR failed for {file_path}") from error
    return completed.stdout


def _extract_text_from_docx(file_path: Path) -> str:
    """Extract paragraph and table text from a .docx file."""
    document = DocxDocument(str(file_path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_text_from_eml(file_path: Path) -> str:
    """Extract headers + plain-text body from a .eml email file."""
    raw = file_path.read_bytes()
    message = message_from_bytes(raw)
    header_lines = [f"{header}: {message.get(header, '')}" for header in ("From", "To", "Subject", "Date")]
    body = _extract_eml_body(message)
    return "\n".join(header_lines) + "\n\n" + body


def _extract_eml_body(message: Message) -> str:
    """Return the first plain-text part of an email message."""
    if message.is_multipart():
        for part in message.walk():
            if part.get_content_type() == "text/plain":
                payload = part.get_payload(decode=True)
                if isinstance(payload, bytes):
                    return payload.decode(part.get_content_charset() or "utf-8", errors="ignore")
        return ""
    payload = message.get_payload(decode=True)
    if isinstance(payload, bytes):
        return payload.decode(message.get_content_charset() or "utf-8", errors="ignore")
    return str(message.get_payload())


def _extract_text_from_xlsx(file_path: Path) -> str:
    """Extract cell text from all sheets in an .xlsx workbook."""
    workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                lines.append(row_text)
    return "\n".join(lines)
