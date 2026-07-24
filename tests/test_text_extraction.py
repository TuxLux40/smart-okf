"""Tests for multi-format text extraction."""

from pathlib import Path

import openpyxl
import pytest
from docx import Document as DocxDocument

from app.exceptions import DocumentIngestError, EncryptedDocumentError
from app.services.text_extraction import extract_text_from_file, is_supported_document


def test_is_supported_document_covers_new_suffixes() -> None:
    for suffix in (".pdf", ".txt", ".docx", ".eml", ".csv", ".xlsx"):
        assert is_supported_document(Path(f"file{suffix}"))


@pytest.mark.parametrize("suffix", [".png", ".jpg", ".jpeg"])
def test_extract_text_from_unreadable_image_raises_clear_ocr_error(tmp_path: Path, suffix: str) -> None:
    image_path = tmp_path / f"scan{suffix}"
    image_path.write_bytes(b"\x89PNG fake bytes")

    with pytest.raises(DocumentIngestError, match="OCR"):
        extract_text_from_file(image_path)


def test_extract_text_from_csv_reads_plain_text(tmp_path: Path) -> None:
    csv_path = tmp_path / "records.csv"
    csv_path.write_text("name,value\nfoo,1\n", encoding="utf-8")

    assert "name,value" in extract_text_from_file(csv_path)


def test_extract_text_from_docx_reads_paragraphs_and_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "letter.docx"
    document = DocxDocument()
    document.add_paragraph("Dear Oliver,")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Invoice"
    table.rows[0].cells[1].text = "42"
    document.save(str(docx_path))

    text = extract_text_from_file(docx_path)

    assert "Dear Oliver," in text
    assert "Invoice" in text
    assert "42" in text


def test_encrypted_ooxml_is_detected_and_skipped(tmp_path: Path) -> None:
    # Encrypted .docx/.xlsx are OLE2 compound files, not ZIPs — detect by the OLE magic header.
    encrypted = tmp_path / "protected.docx"
    encrypted.write_bytes(b"\xd0\xcf\x11\xe0" + b"\x00" * 512)
    with pytest.raises(EncryptedDocumentError, match="password-protected"):
        extract_text_from_file(encrypted)


def test_encrypted_document_error_is_a_document_ingest_error(tmp_path: Path) -> None:
    # Subtype relationship: ingest's broad DocumentIngestError/Exception catch still skips it.
    encrypted = tmp_path / "protected.xlsx"
    encrypted.write_bytes(b"\xd0\xcf\x11\xe0" + b"\x00" * 512)
    with pytest.raises(DocumentIngestError):
        extract_text_from_file(encrypted)


def test_extract_text_from_xlsx_reads_cells(tmp_path: Path) -> None:
    xlsx_path = tmp_path / "budget.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Sheet1"
    sheet.append(["Item", "Cost"])
    sheet.append(["Rent", 900])
    workbook.save(str(xlsx_path))

    text = extract_text_from_file(xlsx_path)

    assert "Sheet1" in text
    assert "Rent" in text
    assert "900" in text


def test_extract_text_from_eml_reads_headers_and_body(tmp_path: Path) -> None:
    eml_path = tmp_path / "message.eml"
    eml_path.write_text(
        "From: doctor@example.com\n"
        "To: oliver@example.com\n"
        "Subject: Appointment reminder\n"
        "Date: Mon, 1 Jun 2026 10:00:00 +0000\n"
        "Content-Type: text/plain; charset=utf-8\n"
        "\n"
        "Your appointment is on Friday.\n",
        encoding="utf-8",
    )

    text = extract_text_from_file(eml_path)

    assert "Subject: Appointment reminder" in text
    assert "Your appointment is on Friday." in text


def test_non_pdf_files_never_touch_marker(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """marker is only ever consulted for .pdf, regardless of ExtractionOptions default."""
    import shutil

    def _fail_if_called(*args: object, **kwargs: object) -> None:
        raise AssertionError("shutil.which must not be called for non-PDF files")

    monkeypatch.setattr(shutil, "which", _fail_if_called)

    text_path = tmp_path / "notes.txt"
    text_path.write_text("plain text, no marker involved", encoding="utf-8")

    assert extract_text_from_file(text_path) == "plain text, no marker involved"


def test_use_marker_defaults_to_true_for_pdf(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """marker is on by default (opt-out): a PDF call with default options still tries it."""
    import shutil

    monkeypatch.setattr(shutil, "which", lambda _name: None)
    pdf_path = tmp_path / "doc.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%%EOF")

    with pytest.raises(DocumentIngestError, match="marker_single"):
        extract_text_from_file(pdf_path)


def test_use_marker_true_raises_clear_error_when_binary_missing(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    import shutil

    from app.services.extraction_options import ExtractionOptions

    monkeypatch.setattr(shutil, "which", lambda _name: None)
    pdf_path = tmp_path / "doc.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%%EOF")

    with pytest.raises(DocumentIngestError, match="marker_single"):
        extract_text_from_file(pdf_path, ExtractionOptions(use_marker=True))


def test_use_marker_true_uses_marker_output_when_binary_present(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    import shutil
    import subprocess
    from pathlib import Path as _Path

    from app.services.extraction_options import ExtractionOptions

    monkeypatch.setattr(shutil, "which", lambda _name: "/usr/local/bin/marker_single")

    def _fake_run(cmd: list[str], **kwargs: object) -> object:
        out_dir = _Path(cmd[cmd.index("--output_dir") + 1])
        (out_dir / "doc.md").write_text("# Marker Output\n\nExtracted table content.", encoding="utf-8")
        return subprocess.CompletedProcess(cmd, 0)

    monkeypatch.setattr(subprocess, "run", _fake_run)
    pdf_path = tmp_path / "doc.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%%EOF")

    text = extract_text_from_file(pdf_path, ExtractionOptions(use_marker=True))

    assert "Extracted table content." in text


def test_read_only_extraction_never_ocr_rewrites_scanned_pdf(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """LIGHT_EXTRACTION (transcript backfill) must raise on a scanned PDF, not mutate it."""
    from app.services import text_extraction
    from app.services.extraction_options import LIGHT_EXTRACTION

    monkeypatch.setattr(text_extraction, "_read_pdf_text", lambda _path: "")

    def _must_not_be_called(_path: Path) -> None:
        raise AssertionError("ocr_pdf_in_place must not run in a read-only pass")

    monkeypatch.setattr(text_extraction, "ocr_pdf_in_place", _must_not_be_called)

    pdf_path = tmp_path / "scan.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%%EOF")

    with pytest.raises(DocumentIngestError, match="read-only"):
        extract_text_from_file(pdf_path, LIGHT_EXTRACTION)
